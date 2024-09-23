import logging
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from telegram import Update, InputFile, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, ContextTypes, ConversationHandler, MessageHandler, filters

# Configuración del logging
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

# Definición de los estados de la conversación
FACTURA_NO, FECHA, DESCRIPCION_CORTA, PRIMER_IMPORTE, IMPORTE_IVA, TOTAL_FACTURA = range(6)

# Función para rellenar la plantilla de Word
def fill_word_template(template_path, output_path, data):
    try:
        doc = Document(template_path)
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f'{{{{{key}}}}}'  # Formato de marcador de posición
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    # Alinear al centro si es descripcion_corta, de lo contrario a la derecha
                    if key == 'descripcion_corta':
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    else:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f'{{{{{key}}}}}'  # Formato de marcador de posición
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
                            # Alinear al centro si es descripcion_corta, de lo contrario a la derecha
                            if key == 'descripcion_corta':
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            else:
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        doc.save(output_path)
    except Exception as e:
        logger.error(f"Error al rellenar la plantilla de Word: {e}")

# Función para iniciar el bot
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text('¡Hola! Vamos a crear una nueva factura. Por favor, proporciona el número de factura:')
    return FACTURA_NO

# Funciones para manejar los estados de la conversación
async def factura_no(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['factura_no'] = update.message.text
    await update.message.reply_text('Gracias. Ahora, por favor, proporciona la fecha de la factura (formato: DD/MM/AAAA):')
    return FECHA

async def fecha(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['fecha'] = update.message.text
    await update.message.reply_text('Por favor, proporciona una descripción corta:')
    return DESCRIPCION_CORTA

async def descripcion_corta(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['descripcion_corta'] = update.message.text
    await update.message.reply_text('Por favor, proporciona el primer importe:')
    return PRIMER_IMPORTE

async def primer_importe(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['primer_importe'] = update.message.text
    await update.message.reply_text('Por favor, proporciona el importe del IVA:')
    return IMPORTE_IVA

async def importe_iva(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['importe_iva'] = update.message.text
    await update.message.reply_text('Por favor, proporciona el total de la factura:')
    return TOTAL_FACTURA

async def total_factura(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['total_factura'] = update.message.text

    invoice_data = {
        'factura_no': context.user_data['factura_no'],
        'fecha': context.user_data['fecha'],
        'descripcion_corta': context.user_data['descripcion_corta'],
        'primer_importe': context.user_data['primer_importe'],
        'importe_iva': context.user_data['importe_iva'],
        'total_factura': context.user_data['total_factura']
    }

    template_path = 'plantilla_factura.docx'
    invoice_file_name = f'invoice_{invoice_data["factura_no"]}.docx'
    fill_word_template(template_path, invoice_file_name, invoice_data)

    if not os.path.exists(invoice_file_name):
        await update.message.reply_text("Error: No se pudo crear el archivo de Word.")
        return ConversationHandler.END

    await update.message.reply_text('Aquí está tu factura en formato Word:')
    with open(invoice_file_name, 'rb') as word_file:
        await update.message.reply_document(document=InputFile(word_file, filename=invoice_file_name))

    # Botón que lleva a la página de conversión
    keyboard = [
        [InlineKeyboardButton("Convertir a PDF", url='https://www.ilovepdf.com/es/word_a_pdf')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text('Descarga el archivo y haz clic en el botón a continuación para convertir tu archivo a PDF:', reply_markup=reply_markup)
    
    # Mensaje adicional
    await update.message.reply_text("Para reiniciar el bot y crear una nueva factura, usa el comando /start")

    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text('Conversación cancelada. Si deseas crear una factura, usa el comando /start.')
    return ConversationHandler.END

def main() -> None:
    application = ApplicationBuilder().token("7554000629:AAHy05nWrWIBnQX3QJqDkrbzaCa4JXS2oqc").build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            FACTURA_NO: [MessageHandler(filters.TEXT & ~filters.COMMAND, factura_no)],
            FECHA: [MessageHandler(filters.TEXT & ~filters.COMMAND, fecha)],
            DESCRIPCION_CORTA: [MessageHandler(filters.TEXT & ~filters.COMMAND, descripcion_corta)],
            PRIMER_IMPORTE: [MessageHandler(filters.TEXT & ~filters.COMMAND, primer_importe)],
            IMPORTE_IVA: [MessageHandler(filters.TEXT & ~filters.COMMAND, importe_iva)],
            TOTAL_FACTURA: [MessageHandler(filters.TEXT & ~filters.COMMAND, total_factura)],
        },
        fallbacks=[CommandHandler('cancel', cancel)],
    )

    application.add_handler(conv_handler)

    # Iniciar el bot
    application.run_polling()

if __name__ == '__main__':
    main()